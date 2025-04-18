import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
import platform
from pathlib import Path
import os

# 获取当前模块所在目录
MODULE_DIR = Path(__file__).parent

# 默认路径配置
DEFAULT_PATHS = {
    "template": MODULE_DIR / "template1.docx",
    "output": MODULE_DIR / "output1.docx",
    "temp_images": MODULE_DIR / "temp_images",
}


def configure_matplotlib_for_chinese():
    system = platform.system()
    if system == "Darwin":  # macOS
        plt.rcParams["font.family"] = ["Arial Unicode MS", "Heiti TC", "sans-serif"]
    elif system == "Windows":
        plt.rcParams["font.family"] = ["Microsoft YaHei", "SimHei", "sans-serif"]
    elif system == "Linux":
        plt.rcParams["font.family"] = ["WenQuanYi Micro Hei", "Droid Sans Fallback", "sans-serif"]
    plt.rcParams["axes.unicode_minus"] = False  # Fix minus sign display issues


configure_matplotlib_for_chinese()


def energy_consumption_by_year(df: pd.DataFrame, energy_type: str) -> pd.Series:
    df = df.copy()
    df["year"] = pd.to_datetime(df["year"]).dt.year
    df = df[df["energy_type"] == energy_type]
    return df.groupby("year")["value"].sum()


def energy_intensity_by_year(gdp_df: pd.DataFrame, energy_df: pd.DataFrame) -> pd.Series:
    gdp_df = gdp_df.copy()
    energy_df = energy_df.copy()
    gdp_df["year"] = pd.to_datetime(gdp_df["year"]).dt.year
    energy_df["year"] = pd.to_datetime(energy_df["year"]).dt.year
    return energy_df.groupby("year")["value"].sum() / gdp_df.groupby("year")["value"].sum()


def energy_consumption_by_industry_and_year(df: pd.DataFrame) -> tuple[pd.Series, pd.Series]:
    df = df.copy()
    df["year"] = pd.to_datetime(df["year"]).dt.year
    return (
        df.groupby(["year", "industry"])["value"].sum(),
        df.groupby(["year", "industry", "energy_type"])["value"].sum(),
    )


def plot_industry_energy_pie(
    df: pd.DataFrame,
    year: int,
    title: str = None,
    figsize: tuple = (10, 10),
    colors: list = None,
    max_categories: int = 6,
    save_path: str = None,
) -> plt.Figure:
    # 获取按行业和年份分组的能源消费数据
    consumption_by_industry_year, _ = energy_consumption_by_industry_and_year(df)

    # 将Series转换为DataFrame并重置索引
    df_consumption = consumption_by_industry_year.reset_index()

    # 过滤特定年份的数据
    df_year = df_consumption[df_consumption["year"] == year]

    if df_year.empty:
        print(f"警告: 未找到{year}年的数据")
        return None

    # 按值排序并提取前N个类别
    df_year = df_year.sort_values(by="value", ascending=False).reset_index(drop=True)

    # 如果行业数量超过max_categories，则将剩余行业归为"其他"
    if len(df_year) > max_categories:
        # 提取前N个类别
        top_categories = df_year.iloc[:max_categories].copy()

        # 计算剩余类别的总和
        others_value = df_year.iloc[max_categories:]["value"].sum()

        # 创建"其他"类别的行
        others_row = pd.DataFrame({"year": [year], "industry": ["其他"], "value": [others_value]})

        # 合并前N个类别和"其他"类别
        df_year = pd.concat([top_categories, others_row], ignore_index=True)

    # 准备饼图数据
    industries = df_year["industry"].tolist()
    values = df_year["value"].tolist()

    # 设置默认颜色方案（如果未提供）
    if colors is None:
        # 使用柔和的颜色方案
        colors = plt.cm.Pastel1(np.linspace(0, 1, len(industries)))

    # 设置默认标题（如果未提供）
    if title is None:
        title = f"{year}年各行业能源消费占比"

    # 创建饼图
    fig, ax = plt.subplots(figsize=figsize)
    wedges, texts, autotexts = ax.pie(
        values,
        labels=industries,
        autopct="%1.1f%%",
        startangle=90,
        colors=colors,
        wedgeprops={"edgecolor": "w", "linewidth": 1},
        textprops={"fontsize": 12},
    )

    # 设置自动文本的属性
    for autotext in autotexts:
        autotext.set_fontsize(10)
        autotext.set_weight("bold")

    # 添加标题
    ax.set_title(title, fontsize=16, pad=20)

    # 确保饼图是圆形的
    ax.axis("equal")

    plt.tight_layout()

    # 保存图表（如果指定）
    if save_path:
        plt.savefig(save_path, dpi=300, bbox_inches="tight")

    return fig


def plot_industry_energy_trends(
    df: pd.DataFrame,
    start_year: int = 2020,
    end_year: int = None,
    title: str = None,
    figsize: tuple = (15, 10),
    colormap: str = "tab10",
    max_categories: int = 7,
    save_path: str = None,
) -> plt.Figure:
    # 获取按行业和年份分组的能源消费数据
    consumption_by_industry_year, _ = energy_consumption_by_industry_and_year(df)

    # 将Series转换为DataFrame并重置索引
    df_consumption = consumption_by_industry_year.reset_index()

    # 过滤年份范围
    if end_year is None:
        end_year = df_consumption["year"].max()

    df_filtered = df_consumption[
        (df_consumption["year"] >= start_year) & (df_consumption["year"] <= end_year)
    ]

    if df_filtered.empty:
        print(f"警告: {start_year}至{end_year}年范围内没有数据")
        return None

    # 处理行业类别数量限制
    # 首先依据行业的总消费量进行排序（忽略"其他"类别）
    df_no_other = df_filtered[df_filtered["industry"] != "其他"]
    industry_totals = df_no_other.groupby("industry")["value"].sum().sort_values(ascending=False)

    # 如果行业数量超过max_categories，则将剩余行业归为"其他"
    if len(industry_totals) > max_categories:
        # 获取前N个最大的行业
        top_industries = industry_totals.index[:max_categories].tolist()

        # 将剩余行业标记为"其他"
        df_filtered.loc[
            ~df_filtered["industry"].isin(top_industries) & (df_filtered["industry"] != "其他"),
            "industry",
        ] = "其他"

        # 对其他类别进行汇总
        df_agg = df_filtered.groupby(["year", "industry"])["value"].sum().reset_index()
        df_filtered = df_agg

    # 排序行业 - 按总量降序排列，将"其他"放在最后
    # 先计算每个行业的总量
    industry_order = df_filtered.groupby("industry")["value"].sum().sort_values(ascending=False)

    # 确保"其他"类别放在最后面（如果存在）
    if "其他" in industry_order.index:
        other_value = industry_order["其他"]
        industry_order = industry_order[industry_order.index != "其他"]
        # 附加到最后
        industry_order = pd.concat([industry_order, pd.Series({"其他": other_value})])

    # 使用排序后的顺序
    industries = industry_order.index.tolist()
    years = sorted(df_filtered["year"].unique())

    # 创建数据透视表以便于绘图
    pivot_data = df_filtered.pivot(index="industry", columns="year", values="value").fillna(0)

    # 确保数据按来自行业排序的industry_order的顺序排列
    pivot_data = pivot_data.reindex(industries)

    # 设置图表尺寸和颜色映射
    fig, ax = plt.subplots(figsize=figsize)
    color_map = plt.cm.get_cmap(colormap)

    # 计算绘图参数
    n_industries = len(industries)
    n_years = len(years)
    bar_width = 0.7 / n_years  # 每个年份柱的宽度 - 设置小一些增加间隔
    industry_width = 1.2  # 每个行业组的宽度 - 增加间隔

    # 创建用于增长率标注的辅助坐标轴
    ax2 = ax.twinx()
    ax2.set_ylim([-50, 50])  # 设置符合增长率的y轴范围，可以根据实际数据调整
    ax2.set_ylabel("年度增长率 (%)", fontsize=14)
    ax2.tick_params(axis="y", colors="gray")
    ax2.yaxis.set_major_formatter(mtick.PercentFormatter(decimals=0))
    ax2.spines["right"].set_color("gray")

    # 创建存储每个行业的增长率数据的字典
    growth_rates = {}

    # 绘制分组柱状图
    for i, industry in enumerate(industries):
        industry_data = pivot_data.loc[industry]
        industry_growth_rates = []  # 存储该行业的增长率
        industry_positions = []  # 存储柱子位置以绘制增长率线

        for j, year in enumerate(years):
            # 计算柱的位置：行业位置 + 柱在行业组内的偏移
            bar_position = i * industry_width + (j - n_years / 2 + 0.5) * bar_width
            value = industry_data.get(year, 0)  # 获取对应年份的值，如果不存在则为0
            industry_positions.append(bar_position)  # 记录位置用于绘制增长率线

            # 计算相对于上一年的变化率（如果不是第一年）
            if j > 0 and industry_data.get(years[j - 1], 0) != 0:
                prev_value = industry_data.get(years[j - 1], 0)
                change_pct = (value - prev_value) / prev_value * 100
                industry_growth_rates.append(change_pct)  # 添加到增长率列表
            else:
                industry_growth_rates.append(None)  # 第一年没有增长率

            # 绘制柱状图
            bar = ax.bar(
                bar_position,
                value,
                width=bar_width * 0.9,  # 稍微缩小柱子宽度增加间隔
                color=color_map(i / n_industries),
                alpha=0.8,
                label=industry if j == 0 else "",
                edgecolor="white",
                linewidth=0.5,
            )

        # 存储该行业的增长率数据
        growth_rates[industry] = {"positions": industry_positions, "rates": industry_growth_rates}

    # 为每个行业绘制增长率线
    for industry, data in growth_rates.items():
        positions = data["positions"]
        rates = data["rates"]

        # 过滤掉第一个点（没有增长率）
        valid_positions = [pos for pos, rate in zip(positions[1:], rates[1:]) if rate is not None]
        valid_rates = [rate for rate in rates[1:] if rate is not None]

        if valid_positions and valid_rates:
            industry_index = industries.index(industry)
            line_color = color_map(industry_index / n_industries)

            # 绘制该行业的增长率线
            ax2.plot(valid_positions, valid_rates, "-o", color=line_color, alpha=0.7, linewidth=1.5)

            # 标注最后一个点的增长率值
            if valid_rates:
                ax2.annotate(
                    f"{valid_rates[-1]:.1f}%",
                    xy=(valid_positions[-1], valid_rates[-1]),
                    xytext=(5, 0),
                    textcoords="offset points",
                    ha="left",
                    va="center",
                    color=line_color,
                    fontsize=8,
                )

    # 设置x轴刻度和标签（两级）
    # 第一级：行业位置 - 考虑行业间隔
    ax.set_xticks([i * industry_width for i in range(n_industries)])
    ax.set_xticklabels(industries, fontsize=12, rotation=30, ha="right")

    # 为每个行业组添加年份标签
    for i, industry in enumerate(industries):
        for j, year in enumerate(years):
            bar_position = i * industry_width + (j - n_years / 2 + 0.5) * bar_width
            # 在柱子上方添加年份标签
            ax.text(
                bar_position,
                pivot_data.loc[industry, year] * 1.02,  # 将文本放在柱子的顶部上方
                str(year),
                ha="center",
                va="bottom",
                fontsize=8,
                color="gray",
            )

    # 设置标题和坐标轴标签
    if title is None:
        title = f"{start_year}至{end_year}年各行业能源消费趋势"
    ax.set_title(title, fontsize=16, pad=20)
    ax.set_ylabel("能源消费量", fontsize=14)

    # 调整轴标签
    ax.set_xlabel("行业类别", fontsize=14)

    # 调整y轴范围，确保显示完整的数据
    ax.set_ylim(bottom=0)  # 设置为从0开始
    # 确保有足够空间显示柱子上方的年份标签
    y_max = pivot_data.max().max() * 1.15  # 增加一点空间
    ax.set_ylim(top=y_max)

    # 添加图例
    handles, labels = ax.get_legend_handles_labels()
    unique_labels = dict(zip(labels, handles))
    ax.legend(
        unique_labels.values(), unique_labels.keys(), title="行业", loc="upper right", fontsize=12
    )

    plt.tight_layout()

    # 保存图表（如果指定）
    if save_path:
        plt.savefig(save_path, dpi=300, bbox_inches="tight")

    return fig


def plot_energy_consumption(
    df: pd.DataFrame,
    energy_type: str,
    title: str = "Energy Consumption by Year",
    figsize: tuple = (12, 8),
    baseline_year: int = 2020,
    save_path: str = None,
) -> plt.Figure:
    # Get the energy consumption data
    consumption = energy_consumption_by_year(df, energy_type)

    # Calculate year-over-year change rate
    yoy_change = consumption.pct_change() * 100

    # Calculate change rate compared to baseline year
    if baseline_year in consumption.index:
        baseline_value = consumption[baseline_year]
        baseline_change = ((consumption - baseline_value) / baseline_value) * 100
    else:
        print(
            f"Warning: Baseline year {baseline_year} not found in data. Using first year as baseline."
        )
        baseline_value = consumption.iloc[0]
        baseline_change = ((consumption - baseline_value) / baseline_value) * 100

    # Create the figure and axes
    fig, ax1 = plt.subplots(figsize=figsize)

    # Plot the bar chart with less saturated colors
    bars = ax1.bar(consumption.index, consumption.values, color="lightsteelblue", alpha=0.8)
    ax1.set_xlabel("年份")
    ax1.set_ylabel("能源消费量")
    ax1.set_title(title)

    # Set x-axis ticks to show all years as integers
    all_years = sorted(list(consumption.index))
    ax1.set_xticks(all_years)
    ax1.set_xticklabels([str(int(year)) for year in all_years])
    ax1.tick_params(axis="x", rotation=45)

    # Create secondary y-axis for change rates
    ax2 = ax1.twinx()
    ax2.set_ylabel("变化率 (%)")
    ax2.yaxis.set_major_formatter(mtick.PercentFormatter(decimals=1))

    # Plot the year-over-year change rate line with less saturated color
    line1 = ax2.plot(
        yoy_change.index,
        yoy_change.values,
        color="lightcoral",
        marker="o",
        linestyle="-",
        label="同比变化率",
    )

    # Plot the baseline year change rate line with less saturated color
    line2 = ax2.plot(
        baseline_change.index,
        baseline_change.values,
        color="lightgreen",
        marker="o",
        linestyle="-",
        label=f"与{baseline_year}年相比变化率",
    )

    # No value labels for cleaner look

    # Combine all legends
    lines = line1 + line2
    labels = [l.get_label() for l in lines]
    ax1.legend(lines, labels, loc="upper left")

    plt.tight_layout()

    # Save the figure if specified
    if save_path:
        plt.savefig(save_path, dpi=300, bbox_inches="tight")

    return fig


def plot_energy_intensity(
    gdp_df: pd.DataFrame,
    energy_df: pd.DataFrame,
    title: str = "Energy Intensity by Year",
    figsize: tuple = (12, 8),
    baseline_year: int = 2020,
    save_path: str = None,
) -> plt.Figure:
    # Get the energy intensity data
    intensity = energy_intensity_by_year(gdp_df, energy_df)

    # Calculate year-over-year change rate
    yoy_change = intensity.pct_change() * 100

    # Calculate change rate compared to baseline year
    if baseline_year in intensity.index:
        baseline_value = intensity[baseline_year]
        baseline_change = ((intensity - baseline_value) / baseline_value) * 100
    else:
        print(
            f"Warning: Baseline year {baseline_year} not found in data. Using first year as baseline."
        )
        baseline_value = intensity.iloc[0]
        baseline_change = ((intensity - baseline_value) / baseline_value) * 100

    # Create the figure and axes
    fig, ax1 = plt.subplots(figsize=figsize)

    # Plot the bar chart with less saturated colors
    bars = ax1.bar(intensity.index, intensity.values, color="bisque", alpha=0.8)
    ax1.set_xlabel("年份")
    ax1.set_ylabel("能源强度")
    ax1.set_title(title)

    # Set x-axis ticks to show all years as integers
    all_years = sorted(list(intensity.index))
    ax1.set_xticks(all_years)
    ax1.set_xticklabels([str(int(year)) for year in all_years])
    ax1.tick_params(axis="x", rotation=45)

    # Create secondary y-axis for change rates
    ax2 = ax1.twinx()
    ax2.set_ylabel("变化率 (%)")
    ax2.yaxis.set_major_formatter(mtick.PercentFormatter(decimals=1))

    # Plot the year-over-year change rate line with less saturated color
    line1 = ax2.plot(
        yoy_change.index,
        yoy_change.values,
        color="lightcoral",
        marker="o",
        linestyle="-",
        label="同比变化率",
    )

    # Plot the baseline year change rate line with less saturated color
    line2 = ax2.plot(
        baseline_change.index,
        baseline_change.values,
        color="lightgreen",
        marker="o",
        linestyle="-",
        label=f"与{baseline_year}年相比变化率",
    )

    # No value labels for cleaner look

    # Combine all legends
    lines = line1 + line2
    labels = [l.get_label() for l in lines]
    ax1.legend(lines, labels, loc="upper right")

    plt.tight_layout()

    # Save the figure if specified
    if save_path:
        plt.savefig(save_path, dpi=300, bbox_inches="tight")

    return fig


def get_docx_placeholder_replacement_values(
    datas: list[pd.DataFrame], year: int, province: str, temp_dir=None
):

    # 创建用于保存结果的字典
    res = {
        "year": year,
        "province": province,
        "values": [],
        "choices": [],
        "industries": [],
        "image_paths": [],
    }

    # 确保临时目录存在
    if temp_dir is None:
        temp_dir = DEFAULT_PATHS["temp_images"]
    os.makedirs(temp_dir, exist_ok=True)
    df1, df2, df3, df4 = datas
    consumption = energy_consumption_by_year(df1, "其他能源")
    intensity = energy_intensity_by_year(df4, df2)
    res["values"].append(consumption[year])  # val1
    res["values"].append(consumption[year - 1])  # val2
    if consumption[year] > consumption[year - 1]:
        res["choices"].append("增长")
    else:
        res["choices"].append("下降")  # choices1
    res["values"].append(consumption[year] / consumption[year - 1])  # val3

    if intensity[year] > intensity[year - 1]:
        res["choices"].append("增长")
    else:
        res["choices"].append("下降")  # choices2
    res["values"].append(intensity[year] / intensity[year - 1])  # val4
    res["values"].append(consumption[2020])  # val5
    res["values"].append(consumption[year])  # val6
    res["values"].append(consumption[year] / consumption[2020])  # val7
    res["values"].append(consumption[year] / consumption[2020] / (year - 2020))  # val8

    ### Trend
    if consumption[year] / consumption[year - 1] > consumption[year] / consumption[2020] / (
        year - 2020
    ):
        res["choices"].append("高于")
    else:
        res["choices"].append("低于")  # choices3
    res["values"].append((intensity[2020] - intensity[year]) / intensity[2020])  # val9
    res["values"].append((intensity[2020] - intensity[year]) / intensity[2020] / 0.135)  # val10
    res["values"].append(year - 2020)  # val11
    res["values"].append(
        (13.5 - (intensity[2020] - intensity[year]) / intensity[2020]) / (year - 2020)
    )  # val12

    consumption_iy1, consumption_iy2 = energy_consumption_by_industry_and_year(df2)
    consumption_iy1 = consumption_iy1.reset_index()
    consumption_iy2 = consumption_iy2.reset_index()

    consumption_y1 = consumption_iy1[consumption_iy1["year"] == year].sort_values(
        by="value", ascending=False
    )

    res["industries"].append(consumption_y1.iloc[0]["industry"])  # industry1
    res["industries"].append(consumption_y1.iloc[1]["industry"])  # industry2

    res["values"].append(0)  # val13
    res["values"].append(0)  # val14

    # 使用 consumption_iy1 计算不同行业的相比前一年的增速
    industry_speedup = {}
    for industry in consumption_iy1["industry"].unique():
        industry_speedup[industry] = (
            consumption_iy1[
                (consumption_iy1["year"] == year) & (consumption_iy1["industry"] == industry)
            ]["value"].values[0]
            / consumption_iy1[
                (consumption_iy1["year"] == year - 1) & (consumption_iy1["industry"] == industry)
            ]["value"].values[0]
        )
    # 排序，获得加速前2的行业和加速比
    sorted_speedup = sorted(industry_speedup.items(), key=lambda x: x[1], reverse=True)
    res["values"].append(sorted_speedup[0][1])  # val15
    res["values"].append(sorted_speedup[1][1])  # val16
    res["industries"].append(sorted_speedup[0][0])  # industry3
    res["industries"].append(sorted_speedup[1][0])  # industry4

    ### Trend
    if (
        consumption_iy1[consumption_iy1["year"] == year]["value"].sum()
        > consumption_iy1[consumption_iy1["year"] == 2020]["value"].sum()
    ):
        res["choices"].append("增长")
    else:
        res["choices"].append("下降")  # choices4

    # 生成并保存图表
    # df1, df2, df3, df4 = datas

    # # 图1: 能源消费量图表
    # fig1 = plot_energy_consumption(
    #     df1,
    #     "其他能源",
    #     title=f"{province}{year}年能源消费量年度变化",
    #     figsize=(10, 6),
    #     baseline_year=2020,
    # )
    # img_path1 = os.path.join(temp_dir, "energy_consumption.png")
    # fig1.savefig(img_path1, dpi=300, bbox_inches="tight")
    # plt.close(fig1)
    # res["image_paths"].append(img_path1)

    # # 图2: 能源强度图表
    # fig2 = plot_energy_intensity(
    #     df4, df2, title=f"{province}{year}年能源强度年度变化", figsize=(10, 6), baseline_year=2020
    # )
    # img_path2 = os.path.join(temp_dir, "energy_intensity.png")
    # fig2.savefig(img_path2, dpi=300, bbox_inches="tight")
    # plt.close(fig2)
    # res["image_paths"].append(img_path2)

    # 图3: 行业能源消费占比饼图
    # fig3 = plot_industry_energy_pie(
    #     df2, year=year, title=f"{province}{year}年各行业能源消费占比", figsize=(8, 8)
    # )
    # img_path3 = os.path.join(temp_dir, "industry_pie.png")
    # fig3.savefig(img_path3, dpi=300, bbox_inches="tight")
    # plt.close(fig3)
    # res["image_paths"].append(img_path3)

    # # 图4: 行业能源消费趋势图
    # fig4 = plot_industry_energy_trends(
    #     df2,
    #     start_year=2020,
    #     end_year=year,
    #     title=f"{province}2020至{year}年各行业能源消费趋势",
    #     figsize=(12, 8),
    #     colormap="Pastel1",
    # )
    # img_path4 = os.path.join(temp_dir, "industry_trends.png")
    # fig4.savefig(img_path4, dpi=300, bbox_inches="tight")
    # plt.close(fig4)
    # res["image_paths"].append(img_path4)

    # print(f"已生成 {len(res['image_paths'])} 张图表用于报告")

    # print(res)

    return res


from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re
import openai
import concurrent.futures
from typing import List, Tuple, Dict, Any


def generate_text_with_llm(context, prompt, max_tokens=150, temperature=0.7):
    """
    使用LLM生成文本的通用函数
    """
    client = openai.OpenAI(
        api_key="sk-esfitgzkgjgnrgjhxpfsuolsgqubdvplclhzpdkuspwpsort",
        base_url="https://api.siliconflow.cn/v1",
    )

    # 使用OpenAI SDK调用API
    response = client.chat.completions.create(
        model="Pro/deepseek-ai/DeepSeek-V3",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens,
        temperature=temperature,
    )

    # 提取生成的文本
    generated_text = response.choices[0].message.content.strip()
    return generated_text


def generate_conclusion_with_llm(context, year, prompt=None):
    """
    生成段落总结
    """
    # 如果没有提供自定义提示词，使用默认提示词
    if prompt is None:
        prompt = f"请基于以下内容，生成简洁的总结，必须以'{year}年'开头\n\n{context}。注意不要重复复述内容中的数据，而应该总结出一些新的观点和见解。"

    return generate_text_with_llm(context, prompt)


def generate_abstract_with_llm(full_text, year, province):
    """
    生成抽象性的文档摘要
    """
    prompt = f"""请基于以下{year}年{province}能源消费分析报告的全文内容，生成一份有逼辨力的摘要。
    
要求：
1. 摘要应讲清楚报告的核心发现与见解
2. 避免重复原文数据，专注于深度的见解和含义
3. 所有观点必须基于原文内容，但要提出独到的分析
4. 摘要的文字量为100-150字
5. 内容在一个段落内，不要添加其他非段落符号

原文内容：
{full_text}
"""

    return generate_text_with_llm(full_text, prompt, max_tokens=500, temperature=0.5)


def replace_docx_placeholders(
    replacement_values: dict[str, any], template_path=None, output_path=None
):
    # 使用默认路径或自定义路径
    if template_path is None:
        template_path = DEFAULT_PATHS["template"]
    if output_path is None:
        output_path = DEFAULT_PATHS["output"]

    # 确保路径是Path对象
    template_path = Path(template_path)
    output_path = Path(output_path)

    print(f"处理Word文档: {template_path}")
    doc = Document(str(template_path))

    if "image_paths" in replacement_values:
        print(f"准备插入{len(replacement_values['image_paths'])}张图片")

    # 替换基本占位符
    for paragraph in doc.paragraphs:
        # 替换年份和省份
        if "<placeholder_year>" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "<placeholder_year>", str(replacement_values["year"])
            )
        if "<placeholder_prev_year>" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "<placeholder_prev_year>", str(replacement_values["year"] - 1)
            )
        if "<placeholder_province>" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "<placeholder_province>", replacement_values["province"]
            )

        # 替换数值
        for i, value in enumerate(replacement_values["values"]):
            if f"<placeholder_val{i+1}>" in paragraph.text:
                # 保留两位小数
                paragraph.text = paragraph.text.replace(
                    f"<placeholder_val{i+1}>", "{:.2f}".format(value)
                )

        # 替换选择项
        for i, choice in enumerate(replacement_values["choices"]):
            if f"<placeholder_choices{i+1}>" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"<placeholder_choices{i+1}>", str(choice))

        # 替换行业名称
        for i, industry in enumerate(replacement_values.get("industries", [])):
            if f"<placeholder_industry{i+1}>" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    f"<placeholder_industry{i+1}>", str(industry)
                )
        # 替换图像
        # for i, img_path in enumerate(replacement_values.get("image_paths", [])):
        #     placeholder = f"<placeholder_img{i+1}>"
        #     if placeholder in paragraph.text:
        #         print(f"处理图片 {i+1}: {placeholder} 在段落: '{paragraph.text[:30]}...'")
        #         replace_placeholder_with_img(paragraph, placeholder, img_path)

    # 处理总结占位符 - 在所有其他占位符替换后执行
    conclusion_placeholder_pattern = r"<placeholder[^>]*conclusion[^>]*>"

    print("开始处理段落总结...")
    year = replacement_values["year"]

    # 使用多线程并发处理所有总结生成任务
    # 步骤1: 收集所有需要处理的段落和上下文
    conclusion_tasks = []
    for i, paragraph in enumerate(doc.paragraphs):
        if re.search(conclusion_placeholder_pattern, paragraph.text):
            original_text = paragraph.text
            print(f"找到总结占位符: '{original_text[:50]}...'")

            # 提取上下文
            context = re.sub(conclusion_placeholder_pattern, "", original_text).strip()

            # 将任务添加到列表
            conclusion_tasks.append((i, paragraph, context))

    # 定义处理单个总结的函数
    def process_conclusion(task_tuple):
        idx, para, ctx = task_tuple
        if ctx.strip():
            # 有内容，生成总结
            print(f"处理总结任务 {idx+1}/{len(conclusion_tasks)}: 使用段落内容生成")
            conclusion = generate_conclusion_with_llm(ctx, year)
        else:
            # 无内容，使用默认总结
            print(f"处理总结任务 {idx+1}/{len(conclusion_tasks)}: 使用默认总结")
            conclusion = f"{year}年能耗数据分析显示能源消费和强度指标有所变化，各行业能源结构存在差异。需要继续关注节能降耗和能源转型发展。"
        return idx, para, conclusion

    # 步骤2: 使用线程池并发处理所有总结生成任务
    if conclusion_tasks:
        print(f"并发处理 {len(conclusion_tasks)} 个总结生成任务...")
        with concurrent.futures.ThreadPoolExecutor(
            max_workers=min(5, len(conclusion_tasks))
        ) as executor:
            # 提交所有任务并等待完成
            futures = {executor.submit(process_conclusion, task): task for task in conclusion_tasks}

            # 处理结果并替换段落内容
            for future in concurrent.futures.as_completed(futures):
                idx, para, conclusion = future.result()
                print(f"完成总结任务 {idx+1}/{len(conclusion_tasks)}: {conclusion[:50]}...")
                # 替换占位符
                para.text = re.sub(conclusion_placeholder_pattern, conclusion, para.text)
    else:
        print("没有找到需要生成总结的段落")

    # 在处理完所有占位符后，生成文档摘要
    abstract_placeholder = "<placeholder_abstract>"

    # 收集文档的全部内容
    print("开始生成文档摘要...")
    full_text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text + "\n\n"

    # 查找摘要占位符
    for paragraph in doc.paragraphs:
        if abstract_placeholder in paragraph.text:
            print(f"找到摘要占位符，生成文档摘要")

            # 生成摘要
            province = replacement_values["province"]
            abstract = generate_abstract_with_llm(full_text, year, province)
            print(f"摘要生成成功: {abstract[:50]}...")

            # 替换占位符
            paragraph.text = paragraph.text.replace(abstract_placeholder, abstract)
            break

    # 保存文档
    output_path.parent.mkdir(parents=True, exist_ok=True)  # 确保输出目录存在
    doc.save(str(output_path))
    print(f"报告已生成: {output_path}")

    # 返回输出文件路径
    return output_path


if __name__ == "__main__":
    df1 = pd.read_excel(
        "/Users/huc/Documents/GitHub_my/intelli-vis/server/data/湖北_外部能耗数据.xlsx",
        sheet_name="hubei_in_y_pro_ind_ene_off",
    )
    df2 = pd.read_excel(
        "/Users/huc/Documents/GitHub_my/intelli-vis/server/data/湖北_外部能耗数据.xlsx",
        sheet_name="hubei_in_y_pro_ind_ene2_off",
    )
    df3 = pd.read_excel(
        "/Users/huc/Documents/GitHub_my/intelli-vis/server/data/湖北_外部能耗数据.xlsx",
        sheet_name="hubei_in_y_pro_ind_prd_off",
    )
    df4 = pd.read_excel(
        "/Users/huc/Documents/GitHub_my/intelli-vis/server/data/湖北_外部能耗数据.xlsx",
        sheet_name="hubei_in_y_pro_gdp_off",
    )

    # print(energy_consumption_by_year(df1, "其他能源"))
    # print(energy_intensity_by_year(df4, df2))

    # plot_energy_consumption(df1, "其他能源", title="其他能源消费量年度变化")
    # plot_energy_intensity(df4, df2, title="能源强度年度变化")
    # plt.show()

    # print(energy_consumption_by_industry_and_year(df2))
    # plot_industry_energy_pie(df2, year=2020)

    # plot_industry_energy_trends(df2, start_year=2020, end_year=2022)
    # plt.show()

    replacement_values = get_docx_placeholder_replacement_values(
        [df1, df2, df3, df4], year=2022, province="湖北"
    )
    replace_docx_placeholders(replacement_values)
